from __future__ import annotations

import re
from copy import copy
from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document


ROOT = Path(__file__).resolve().parent.parent
HU_DIR = ROOT / "insumos_cp" / "hu"
CHECKLIST_PATH = ROOT / "insumos_cp" / "Checklist de Historia de Usuario para generación de casos de prueba.docx"
FORMAT_JIRA_PATH = ROOT / "insumos_cp" / "Formato Jira.xlsx"
OUTPUT_PATH = ROOT / "casos de prueba" / "casos everest.xlsx"


def _read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_hu_files() -> list[Path]:
    return sorted(HU_DIR.glob("*.md"))


def _slug_to_hu_id(name: str) -> str:
    m = re.match(r"(HU-\d+)", name)
    return m.group(1) if m else "HU-000"


def _parse_hu(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="ignore")
    hu_id = _slug_to_hu_id(path.name)
    title = path.stem.replace("-", " ").replace("_", " ").strip()

    # Heurísticas ligeras para enriquecer los casos sin inventar lógica de negocio
    words = title.lower()
    if "consulta general" in words:
        domain = "consulta general"
    elif "cartera detallada" in words:
        domain = "cartera detallada"
    elif "tc detallada" in words:
        domain = "tarjeta de crédito"
    elif "cdt detallado" in words:
        domain = "cdt detallado"
    else:
        domain = title

    return {
        "id": hu_id,
        "title": title,
        "domain": domain,
        "raw": text,
    }


def _build_cases_for_hu(hu: dict) -> list[dict]:
    hu_id = hu["id"]
    title = hu["title"]
    domain = hu["domain"]

    resumen_base = f"[{hu_id}] {title}"

    return [
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - flujo feliz principal",
            "Descripcion": f"Validar el flujo principal y la consolidación correcta de {domain}.",
            "Escenario": f"HU {hu_id} cargada con datos válidos y servicios dependientes disponibles.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo principal definido en la HU.",
            "Datos": "Datos válidos de entrada según la historia de usuario.",
            "Resultado Esperado": "Respuesta consolidada exitosa con la información solicitada.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de campo obligatorio",
            "Descripcion": f"Validar el comportamiento ante un campo obligatorio ausente en {domain}.",
            "Escenario": f"HU {hu_id} con un campo obligatorio ausente.",
            "Resultado Final": "Pending",
            "Accion": "Enviar la solicitud omitiendo un campo obligatorio.",
            "Datos": "Body incompleto según la HU.",
            "Resultado Esperado": "El sistema rechaza la solicitud y muestra la validación esperada.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de formato inválido",
            "Descripcion": f"Validar el comportamiento ante un dato con formato inválido en {domain}.",
            "Escenario": f"HU {hu_id} con un dato en formato incorrecto.",
            "Resultado Final": "Pending",
            "Accion": "Enviar la solicitud con un valor de formato inválido.",
            "Datos": "Dato con formato no permitido por la HU.",
            "Resultado Esperado": "El sistema informa el error esperado y no procesa la solicitud.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de valor límite inferior",
            "Descripcion": f"Validar el comportamiento con el valor mínimo o límite inferior aplicable a {domain}.",
            "Escenario": f"HU {hu_id} ejecutada con el mínimo permitido o valor límite inferior.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo con el valor mínimo permitido.",
            "Datos": "Valor mínimo o límite inferior según aplique.",
            "Resultado Esperado": "El sistema procesa correctamente el valor mínimo o lo rechaza según la regla de negocio.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de valor límite superior",
            "Descripcion": f"Validar el comportamiento con el valor máximo o límite superior aplicable a {domain}.",
            "Escenario": f"HU {hu_id} ejecutada con el máximo permitido o valor límite superior.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo con el valor máximo permitido.",
            "Datos": "Valor máximo o límite superior según aplique.",
            "Resultado Esperado": "El sistema procesa correctamente el valor máximo o lo rechaza según la regla de negocio.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de dato nulo",
            "Descripcion": f"Validar el comportamiento cuando un campo llega nulo en {domain}.",
            "Escenario": f"HU {hu_id} con un campo nulo.",
            "Resultado Final": "Pending",
            "Accion": "Enviar la solicitud con un valor nulo en un campo relevante.",
            "Datos": "Campo con valor null.",
            "Resultado Esperado": "El sistema rechaza el dato nulo o lo trata según la regla definida.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - validación de dato vacío",
            "Descripcion": f"Validar el comportamiento cuando un campo llega vacío en {domain}.",
            "Escenario": f"HU {hu_id} con un campo vacío.",
            "Resultado Final": "Pending",
            "Accion": "Enviar la solicitud con un valor vacío en un campo relevante.",
            "Datos": "Campo vacío.",
            "Resultado Esperado": "El sistema rechaza el dato vacío o lo trata según la regla definida.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - combinación de campos obligatorios",
            "Descripcion": f"Validar la combinación de varios campos obligatorios para {domain}.",
            "Escenario": f"HU {hu_id} con más de un campo obligatorio en combinación.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo con la combinación de campos obligatorios.",
            "Datos": "Combinación válida o inválida según la HU.",
            "Resultado Esperado": "La respuesta cumple la validación o consolidación esperada.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - manejo de error de dependencia",
            "Descripcion": f"Validar que {title} maneja errores de servicios o dependencias externas.",
            "Escenario": f"HU {hu_id} con dependencia externa no disponible o con error.",
            "Resultado Final": "Pending",
            "Accion": "Forzar falla en la dependencia y observar la respuesta del flujo.",
            "Datos": "Simulación de error técnico o funcional de la dependencia.",
            "Resultado Esperado": "Se retorna el error o la respuesta parcial definida por la historia.",
        },
        {
            "Tipo de test": "Funcional",
            "Resumen": f"{resumen_base} - manejo de error parcial",
            "Descripcion": f"Validar el comportamiento ante una respuesta parcial en {domain}.",
            "Escenario": f"HU {hu_id} con una dependencia exitosa y otra fallida.",
            "Resultado Final": "Pending",
            "Accion": "Simular una respuesta parcial y verificar el resultado.",
            "Datos": "Una dependencia disponible y otra con error.",
            "Resultado Esperado": "El sistema responde con el comportamiento parcial definido por la HU.",
        },
        {
            "Tipo de test": "Performance",
            "Resumen": f"{resumen_base} - tiempo de respuesta nominal",
            "Descripcion": f"Validar el tiempo de respuesta para {domain} bajo condiciones normales.",
            "Escenario": f"HU {hu_id} ejecutada en ambiente controlado con datos válidos.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo y medir latencia de extremo a extremo.",
            "Datos": "Carga normal definida por la HU o por el entorno.",
            "Resultado Esperado": "El tiempo de respuesta cumple el umbral definido o queda documentado como pendiente de umbral.",
        },
        {
            "Tipo de test": "Performance",
            "Resumen": f"{resumen_base} - carga repetida",
            "Descripcion": f"Validar el comportamiento de {domain} bajo ejecución repetida.",
            "Escenario": f"HU {hu_id} ejecutada varias veces consecutivas.",
            "Resultado Final": "Pending",
            "Accion": "Repetir el flujo varias veces y observar degradación.",
            "Datos": "Múltiples ejecuciones seguidas.",
            "Resultado Esperado": "No se evidencia degradación funcional y la latencia se mantiene dentro del rango esperado.",
        },
        {
            "Tipo de test": "Performance",
            "Resumen": f"{resumen_base} - volumen de datos",
            "Descripcion": f"Validar el desempeño de {domain} con mayor volumen de datos.",
            "Escenario": f"HU {hu_id} con un set de datos amplio.",
            "Resultado Final": "Pending",
            "Accion": "Ejecutar el flujo con volumen incrementado.",
            "Datos": "Volumen alto de registros o elementos aplicables.",
            "Resultado Esperado": "El sistema mantiene estabilidad y responde dentro del umbral esperado.",
        },
        {
            "Tipo de test": "Performance",
            "Resumen": f"{resumen_base} - concurrencia básica",
            "Descripcion": f"Validar el comportamiento de {domain} con dos o más ejecuciones concurrentes.",
            "Escenario": f"HU {hu_id} con ejecución simultánea de solicitudes.",
            "Resultado Final": "Pending",
            "Accion": "Disparar solicitudes concurrentes y comparar resultados.",
            "Datos": "Solicitudes simultáneas.",
            "Resultado Esperado": "La concurrencia no rompe la consistencia ni la estabilidad del flujo.",
        },
        {
            "Tipo de test": "Performance",
            "Resumen": f"{resumen_base} - tolerancia a picos",
            "Descripcion": f"Validar el comportamiento de {domain} ante un incremento brusco de uso.",
            "Escenario": f"HU {hu_id} con un pico de ejecuciones.",
            "Resultado Final": "Pending",
            "Accion": "Generar un pico de solicitudes sobre el flujo.",
            "Datos": "Incremento súbito de carga.",
            "Resultado Esperado": "El sistema se mantiene operativo o degrada de forma controlada.",
        },
        {
            "Tipo de test": "Accesibilidad",
            "Resumen": f"{resumen_base} - accesibilidad de información",
            "Descripcion": f"Validar que la información expuesta por {title} sea comprensible y trazable para consumo humano o documental.",
            "Escenario": f"HU {hu_id} revisada desde perspectiva de accesibilidad de contenido.",
            "Resultado Final": "Pending",
            "Accion": "Revisar etiquetas, mensajes, estructura de salida o campos visibles según aplique.",
            "Datos": "Contenido generado por el flujo de la HU.",
            "Resultado Esperado": "La información es clara, trazable y usable conforme a la historia y sus reglas.",
        },
        {
            "Tipo de test": "Accesibilidad",
            "Resumen": f"{resumen_base} - consistencia de mensajes",
            "Descripcion": f"Validar la claridad y consistencia de mensajes asociados a {domain}.",
            "Escenario": f"HU {hu_id} con mensajes visibles para el usuario o consumidor.",
            "Resultado Final": "Pending",
            "Accion": "Revisar el texto de validaciones, errores o confirmaciones.",
            "Datos": "Mensajes generados por el flujo.",
            "Resultado Esperado": "Los mensajes son claros, consistentes y comprensibles.",
        },
        {
            "Tipo de test": "Accesibilidad",
            "Resumen": f"{resumen_base} - trazabilidad de campos",
            "Descripcion": f"Validar que los campos y resultados de {domain} sean interpretables sin ambigüedad.",
            "Escenario": f"HU {hu_id} con información estructurada para revisión.",
            "Resultado Final": "Pending",
            "Accion": "Verificar el orden, rotulado y relación entre datos.",
            "Datos": "Campos y resultados del flujo.",
            "Resultado Esperado": "La salida permite identificar el significado de cada campo sin ambigüedad.",
        },
        {
            "Tipo de test": "Accesibilidad",
            "Resumen": f"{resumen_base} - legibilidad de salida",
            "Descripcion": f"Validar la legibilidad de la salida o respuesta de {title}.",
            "Escenario": f"HU {hu_id} con salida visible o documentada.",
            "Resultado Final": "Pending",
            "Accion": "Revisar formato, separación y claridad del contenido.",
            "Datos": "Salida del caso de prueba.",
            "Resultado Esperado": "La salida es legible, ordenada y fácil de interpretar.",
        },
    ]


def _ensure_template():
    if not FORMAT_JIRA_PATH.exists():
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "TestCases"
        headers = [
            "Issue Key",
            "Summary",
            "Description",
            "Precondition",
            "Status",
            "Priority",
            "Assignee",
            "Reporter",
            "Estimated Time",
            "Labels",
            "Components",
            "Sprint",
            "Fix Versions",
            "Is Shareable Step",
            "Shareable Testcase Issue Key",
            "Shareable Testcase Version No.",
            "Step Summary",
            "Test Data",
            "Expected Result",
            "Version",
            "Folders",
            "TestCase Type",
            "Created By",
            "Created Date",
            "Updated By",
            "Updated Date",
            "Story Linkages",
            "Comment Count",
            "Attachment Count",
            "Story Count",
        ]
        ws.append(headers)
        wb.save(FORMAT_JIRA_PATH)


def _copy_style(src, dst):
    if src.has_style:
        dst._style = copy(src._style)
    if src.font:
        dst.font = copy(src.font)
    if src.fill:
        dst.fill = copy(src.fill)
    if src.border:
        dst.border = copy(src.border)
    if src.alignment:
        dst.alignment = copy(src.alignment)
    if src.number_format:
        dst.number_format = src.number_format
    if src.protection:
        dst.protection = copy(src.protection)


def build_excel():
    _ensure_template()
    huns = [_parse_hu(p) for p in _load_hu_files()]
    if not huns:
        raise FileNotFoundError("No se encontraron HUs en insumos_cp/hu")

    checklist_text = _read_docx_text(CHECKLIST_PATH) if CHECKLIST_PATH.exists() else ""
    wb = openpyxl.load_workbook(FORMAT_JIRA_PATH)
    ws = wb.active

    # Limpiar filas de datos previas si existieran
    if ws.max_row > 1:
        ws.delete_rows(2, ws.max_row - 1)

    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    header_map = {str(name).strip(): idx for idx, name in enumerate(headers, start=1) if name}

    required_headers = [
        "Issue Key",
        "Summary",
        "Description",
        "Precondition",
        "Status",
        "Priority",
        "Labels",
        "Components",
        "Step Summary",
        "Test Data",
        "Expected Result",
        "TestCase Type",
        "Story Linkages",
    ]
    missing = [h for h in required_headers if h not in header_map]
    if missing:
        raise ValueError(f"Faltan columnas en Formato Jira.xlsx: {', '.join(missing)}")

    row_idx = 2
    issue_id = 1

    for hu in huns:
        cases = _build_cases_for_hu(hu)
        for case in cases:
            values = {
                "Issue Key": f"EV-{issue_id}",
                "Summary": case["Resumen"],
                "Description": case["Descripcion"],
                "Precondition": case["Escenario"],
                "Status": "To Do",
                "Priority": "Medium" if case["Tipo de test"] != "Performance" else "High",
                "Assignee": "",
                "Reporter": "",
                "Estimated Time": "",
                "Labels": f"everest,{case['Tipo de test'].lower()}",
                "Components": hu["title"],
                "Sprint": "",
                "Fix Versions": "",
                "Is Shareable Step": "No",
                "Shareable Testcase Issue Key": "",
            "Shareable Testcase Version No.": "",
            "Step Summary": f"1. Preparar el entorno y los datos para el caso: {case['Accion']} | 2. Ejecutar la solicitud o acción principal | 3. Validar la respuesta obtenida contra el resultado esperado",
            "Test Data": case["Datos"],
                "Expected Result": case["Resultado Esperado"],
                "Version": "1",
                "Folders": hu["id"],
                "TestCase Type": case["Tipo de test"],
                "Created By": "",
                "Created Date": "",
                "Updated By": "",
                "Updated Date": "",
                "Story Linkages": hu["id"],
                "Comment Count": "0",
                "Attachment Count": "0",
                "Story Count": "1",
            }

            for name, col_idx in header_map.items():
                ws.cell(row=row_idx, column=col_idx, value=values.get(name, ""))

            row_idx += 1
            issue_id += 1

    for c in range(1, ws.max_column + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    wb.save(OUTPUT_PATH)
    return len(huns), issue_id - 1, checklist_text


if __name__ == "__main__":
    total_hus, total_cases, _ = build_excel()
    print(f"HUs procesadas: {total_hus}")
    print(f"Casos generados: {total_cases}")
    print(f"Archivo generado: {OUTPUT_PATH}")
    print(f"Formato base usado: {FORMAT_JIRA_PATH}")
